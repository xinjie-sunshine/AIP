# -*- coding: utf-8 -*-
# @Author  : xinjie
from ast import Pass
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from datetime import datetime, timedelta
from os import listdir
import os

def toreporter():
    Pass


def toword():
    document = Document()
	#
	# document.add_heading('聚合支付平台主机CPU及内存使用情况', 0)

    h = document.add_heading("聚合支付平台资源使用情况", 1)
    h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.add_run(datetime.now().strftime('%m-%d')+ '日')
    if (str(datetime.now().time())>"16:00"):
        ctime = "日15:00"
    else:
        ctime = "日15:00"
    p = document.add_paragraph("（"+ (datetime.now() - timedelta(days=2)).strftime('%Y-%m-%d')  + ctime + "至" + datetime.now( ).strftime("%Y-%m-%d")+ ctime  + ")")
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if listdir("pic"):            
        document.add_paragraph("一、聚合支付平台paas资源情况")
        p = document.add_paragraph('1、K8S-NS-19资源（cpu/内存使用率）')
        pic = [fn for fn in listdir("pic") if fn.startswith("ns-19")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))

        p = document.add_paragraph('2、好码齐基础库')
        pic = [fn for fn in listdir("pic") if fn.startswith("hmq_mysql_jichu")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))
        p = document.add_paragraph('3、好码齐订单库')
        pic = [fn for fn in listdir("pic") if fn.startswith("hmq_mysql_dingdan")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))

        p = document.add_paragraph('4、es集群状态')
        pic = [fn for fn in listdir("pic") if fn.startswith("es")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))
        
        p = document.add_paragraph('5、组件代理nginx状态')
        pic = [fn for fn in listdir("pic") if fn.startswith("proxy_ng")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))

        p = document.add_paragraph('6、交易情况')
        pic = [fn for fn in listdir("pic") if fn.startswith("kefen")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))

        p = document.add_paragraph('7、听云APM')
        pic = [fn for fn in listdir("pic") if fn.startswith("tingyun")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))
  


        # 商盟
        h = document.add_heading("商户联盟及周边系统底层维护巡检结果", 1)
        h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = document.add_paragraph('1、K8S-NS-99资源（cpu/内存使用率）')
        pic = [fn for fn in listdir("pic") if fn.startswith("ns-99")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))

        p = document.add_paragraph('2.	数据库状态：')
        pic = [fn for fn in listdir("pic") if fn.startswith("shlm_mysql")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))
        p = document.add_paragraph('2.	 Redis状态:')
        pic = [fn for fn in listdir("pic") if fn.startswith("shlm_redis")]
        pic.sort()
        for fn in pic:
            document.add_picture("pic/" + fn, width=Inches(6.25))
            
        document.save(datetime.now().strftime("%m-%d") + '聚合支付平台资源使用情况.docx')
		# 删除保存的图片
        pic = [fn for fn in listdir("pic") if fn.endswith(".png")]
        for fn in pic:
            os.remove("pic/" + fn)
    else:
        print("未获取到图片")


if __name__ == "__main__":
    toword()

