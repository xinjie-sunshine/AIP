# -*- coding: utf-8 -*-
# @Author  : xinjie
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from datetime import datetime, timedelta
from os import listdir
import os
from excel_ops import ExcelUtil
from utils.logger import Logger

# create a logger instance
logger = Logger(logger='BasePage').getlog()


def toword():


    # 获取巡检对象
    filepath = "AIP.xlsx"
    sheetName = "pic_reporter_prefix"    
    data = ExcelUtil(filepath, sheetName)
    obs_list = data.dict_data()
    main_title = data.get_main_title()
    pic_deleted = data.get_pic_delete()

    document = Document()

    logger.info("生成Word报告标题为%s" % main_title)
    ## 标题及时间
    h = document.add_heading(main_title, 1)
    h.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # h.add_run(datetime.now().strftime('%m-%d')+ '日')
    p = document.add_paragraph("巡检时间："+ (datetime.now()).strftime('%Y-%m-%d %H:%M:%S')  )
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if listdir("pic"):
        for item in obs_list:
            p = document.add_paragraph(item["title"])
            pic = [fn for fn in listdir("pic") if fn.startswith(item["prefix"])]
            # pic.sort()
            for fn in pic:
                document.add_picture("pic/" + fn, width=Inches(6.25))
            
        document.save(datetime.now().strftime("%m-%d") + main_title +'.docx')
		# 删除保存的图片
        if pic_deleted == "Yes":
            logger.info("pic目录将会清空")
            pic = [fn for fn in listdir("pic") if fn.endswith(".png")]
            for fn in pic:
                os.remove("pic/" + fn)
    else:
        print("未获取到图片")


if __name__ == "__main__":
    toword()

